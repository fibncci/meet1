from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, abort, send_from_directory, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta, time
from functools import wraps
import calendar
import json
import os
import uuid
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
############ 连接数据库  Flask 怎么连这个 SQLite 的？
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///meeting_rooms.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# 文件上传配置
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 最大10MB
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx', 'ppt', 'pptx', 'xls', 'xlsx', 'txt'}

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 添加uploads目录到.gitignore，避免上传文件纳入版本控制
gitignore_path = os.path.join(os.path.dirname(app.root_path), '.gitignore')
if os.path.exists(gitignore_path):
    with open(gitignore_path, 'r') as f:
        gitignore_content = f.read()
    
    if 'uploads/' not in gitignore_content:
        with open(gitignore_path, 'a') as f:
            f.write('\n# 会议文档上传目录\nuploads/\n')

# 判断文件扩展名是否允许上传
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

### 连接数据库  Flask 怎么连这个 SQLite 的？
db = SQLAlchemy(app)
########### app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///instance/meeting_rooms.db'
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = '请先登录'
login_manager.login_message_category = 'info'

# 用户角色
class Role(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(20), unique=True, nullable=False)
    users = db.relationship('User', backref='role', lazy=True)

# 用户模型
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    department = db.Column(db.String(100))
    phone = db.Column(db.String(20))
    role_id = db.Column(db.Integer, db.ForeignKey('role.id'), nullable=False)
    reservations = db.relationship('Reservation', backref='user', lazy=True)
    uploaded_documents = db.relationship('ReservationDocument', back_populates='uploader', lazy=True, foreign_keys='ReservationDocument.uploaded_by')
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
        
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    def is_admin(self):
        return self.role.name == 'admin'

# 会议室设备
class Equipment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    
    room_equipments = db.relationship('RoomEquipment', backref='equipment', lazy=True)

# 会议室模型
class Room(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    location = db.Column(db.String(100), nullable=False)
    capacity = db.Column(db.Integer, nullable=False)
    description = db.Column(db.Text)
    is_active = db.Column(db.Boolean, default=True)
    reservations = db.relationship('Reservation', backref='room', lazy=True)
    equipments = db.relationship('RoomEquipment', backref='room', lazy=True)
    
    def get_equipment_list(self):
        return [re.equipment for re in self.equipments]

# 会议室设备关联
class RoomEquipment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    room_id = db.Column(db.Integer, db.ForeignKey('room.id'), nullable=False)
    equipment_id = db.Column(db.Integer, db.ForeignKey('equipment.id'), nullable=False)
    quantity = db.Column(db.Integer, default=1)

# 预定模型
class Reservation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    date = db.Column(db.Date, nullable=False)
    start_time = db.Column(db.Time, nullable=False)
    end_time = db.Column(db.Time, nullable=False)
    attendees = db.Column(db.Integer, nullable=False)
    description = db.Column(db.Text)
    status = db.Column(db.String(20), default='confirmed')  # confirmed, canceled, completed
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    room_id = db.Column(db.Integer, db.ForeignKey('room.id'), nullable=False)
    documents = db.relationship('ReservationDocument', backref='reservation', lazy=True, cascade="all, delete-orphan")
    
    def is_past(self):
        now = datetime.now()
        reservation_end = datetime.combine(self.date, self.end_time)
        return now > reservation_end
    
    def can_cancel(self):
        now = datetime.now()
        reservation_start = datetime.combine(self.date, self.start_time)
        return reservation_start > now and self.status == 'confirmed'
    
    def format_time(self):
        return f"{self.start_time.strftime('%H:%M')} - {self.end_time.strftime('%H:%M')}"

# 会议文档模型
class ReservationDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)  # 原始文件名
    stored_filename = db.Column(db.String(255), nullable=False)  # 存储的文件名
    file_type = db.Column(db.String(50))  # 文件类型
    file_size = db.Column(db.Integer)  # 文件大小（字节）
    upload_time = db.Column(db.DateTime, default=datetime.utcnow)
    reservation_id = db.Column(db.Integer, db.ForeignKey('reservation.id'), nullable=False)
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    uploader = db.relationship('User', back_populates='uploaded_documents', foreign_keys=[uploaded_by])

# 会议室维护记录
class Maintenance(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    room_id = db.Column(db.Integer, db.ForeignKey('room.id'), nullable=False)
    start_date = db.Column(db.Date, nullable=False)
    end_date = db.Column(db.Date, nullable=False)
    reason = db.Column(db.Text, nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    room = db.relationship('Room', backref='maintenances')
    creator = db.relationship('User')

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# 管理员权限检查装饰器
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin():
            flash('需要管理员权限', 'danger')
            return redirect(url_for('home'))
        return f(*args, **kwargs)
    return decorated_function

# 初始化数据库
with app.app_context():
    db.create_all()
    
    # 创建角色
    if not Role.query.first():
        roles = [
            Role(name='admin'),
            Role(name='user')
        ]
        db.session.add_all(roles)
        db.session.commit()
    
    # 创建管理员账户
    admin_role = Role.query.filter_by(name='admin').first()
    if not User.query.filter_by(username='admin').first():
        admin = User(
            username='admin',
            email='admin@example.com',
            department='管理部',
            phone='14751366061',
            role_id=admin_role.id
        )
        admin.set_password('admin')
        db.session.add(admin)
        db.session.commit()
    
    # 创建设备类型
    if not Equipment.query.first():
        equipments = [
            Equipment(name='投影仪'),
            Equipment(name='电子白板'),
            Equipment(name='视频会议系统'),
            Equipment(name='音响系统'),
            Equipment(name='电视'),
            Equipment(name='空调')
        ]
        db.session.add_all(equipments)
        db.session.commit()
    
    # 创建默认会议室
    # ### 中软
    # Room(name="会议室B", location="7楼", capacity=20, description="中型会议室"),
    # Room(name="会议室C", location="3楼", capacity=5, description="小型洽谈室"),
    if not Room.query.first():
        rooms = [
            Room(name="会议室A", location="6楼", capacity=19, description="中型会议室"),
            Room(name="会议室D", location="4楼", capacity=50, description="大型会议室")
        ]
        db.session.add_all(rooms)
        db.session.commit()
        
        # 添加会议室设备
        projector = Equipment.query.filter_by(name='投影仪').first()
        whiteboard = Equipment.query.filter_by(name='电子白板').first()
        video_conf = Equipment.query.filter_by(name='视频会议系统').first()
        audio = Equipment.query.filter_by(name='音响系统').first()
        tv = Equipment.query.filter_by(name='电视').first()
        ac = Equipment.query.filter_by(name='空调').first()
        
        room_equipments = [
            RoomEquipment(room_id=1, equipment_id=projector.id),
            RoomEquipment(room_id=1, equipment_id=whiteboard.id),
            RoomEquipment(room_id=1, equipment_id=ac.id),
            
            RoomEquipment(room_id=2, equipment_id=projector.id),
            RoomEquipment(room_id=2, equipment_id=whiteboard.id),
            RoomEquipment(room_id=2, equipment_id=video_conf.id),
            RoomEquipment(room_id=2, equipment_id=audio.id),
            RoomEquipment(room_id=2, equipment_id=ac.id),
            
            RoomEquipment(room_id=3, equipment_id=tv.id),
            RoomEquipment(room_id=3, equipment_id=ac.id),
            
            RoomEquipment(room_id=4, equipment_id=projector.id),
            RoomEquipment(room_id=4, equipment_id=whiteboard.id),
            RoomEquipment(room_id=4, equipment_id=video_conf.id),
            RoomEquipment(room_id=4, equipment_id=audio.id),
            RoomEquipment(room_id=4, equipment_id=ac.id)
        ]
        db.session.add_all(room_equipments)
        db.session.commit()

# 路由
@app.route('/')
def home():
    today = datetime.now().date()
    rooms = Room.query.filter_by(is_active=True).all()
    upcoming_reservations = Reservation.query.filter(
        Reservation.date >= today,
        Reservation.status == 'confirmed'
    ).order_by(Reservation.date, Reservation.start_time).limit(5).all()
    
    return render_template('home.html', 
                          rooms=rooms, 
                          upcoming_reservations=upcoming_reservations,
                          today=today)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
        
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        department = request.form['department']
        phone = request.form['phone']
        
        if User.query.filter_by(username=username).first():
            flash('用户名已存在', 'danger')
            return redirect(url_for('register'))
        
        if User.query.filter_by(email=email).first():
            flash('邮箱已被注册', 'danger')
            return redirect(url_for('register'))
        
        user_role = Role.query.filter_by(name='user').first()
        user = User(
            username=username, 
            email=email, 
            department=department,
            phone=phone,
            role_id=user_role.id
        )
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        flash('注册成功，请登录', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
        
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        remember = 'remember' in request.form
        
        user = User.query.filter_by(username=username).first()
        
        if not user or not user.check_password(password):
            flash('用户名或密码错误', 'danger')
            return redirect(url_for('login'))
        
        login_user(user, remember=remember)
        next_page = request.args.get('next')
        
        flash('登录成功', 'success')
        if next_page:
            return redirect(next_page)
        return redirect(url_for('home'))
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('已退出登录', 'success')
    return redirect(url_for('home'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        email = request.form['email']
        department = request.form['department']
        phone = request.form['phone']
        current_password = request.form['current_password']
        new_password = request.form['new_password']
        
        if email != current_user.email and User.query.filter_by(email=email).first():
            flash('邮箱已被使用', 'danger')
            return redirect(url_for('profile'))
        
        if current_password and new_password:
            if not current_user.check_password(current_password):
                flash('当前密码不正确', 'danger')
                return redirect(url_for('profile'))
            current_user.set_password(new_password)
            flash('密码已更新', 'success')
        
        current_user.email = email
        current_user.department = department
        current_user.phone = phone
        
        db.session.commit()
        flash('个人资料已更新', 'success')
        return redirect(url_for('profile'))
    
    return render_template('profile.html')

@app.route('/rooms')
def room_list():
    rooms = Room.query.filter_by(is_active=True).all()
    return render_template('room_list.html', rooms=rooms)

@app.route('/room/<int:id>')
def room_detail(id):
    room = Room.query.get_or_404(id)
    today = datetime.now().date()
    
    # 获取会议室当天预订
    reservations = Reservation.query.filter_by(
        room_id=id, 
        date=today,
        status='confirmed'
    ).order_by(Reservation.start_time).all()
    
    # 获取维护记录
    maintenances = Maintenance.query.filter_by(room_id=id).filter(
        Maintenance.end_date >= today
    ).order_by(Maintenance.start_date).all()
    
    return render_template('room_detail.html', 
                          room=room, 
                          reservations=reservations,
                          maintenances=maintenances,
                          today=today)

@app.route('/reserve', methods=['GET', 'POST'])
@login_required
def reserve():
    if request.method == 'POST':
        title = request.form['title']
        room_id = request.form['room_id']
        date_str = request.form['date']
        start_time_str = request.form['start_time']
        end_time_str = request.form['end_time']
        attendees = int(request.form['attendees'])
        description = request.form['description']
        
        date = datetime.strptime(date_str, '%Y-%m-%d').date()
        start_time = datetime.strptime(start_time_str, '%H:%M').time()
        end_time = datetime.strptime(end_time_str, '%H:%M').time()
        
        # 检查日期是否有效
        if date < datetime.now().date():
            flash('无法预订过去的日期', 'danger')
            return redirect(url_for('reserve'))
        
        # 检查时间是否有效
        if start_time >= end_time:
            flash('结束时间必须晚于开始时间', 'danger')
            return redirect(url_for('reserve'))
        
        # 检查工作时间（假设8:00-20:00为工作时间）
        work_start = time(8, 0)
        work_end = time(20, 0)
        if start_time < work_start or end_time > work_end:
            flash('预订时间必须在工作时间内（8:00-20:00）', 'danger')
            return redirect(url_for('reserve'))
        
        room = Room.query.get_or_404(room_id)
        
        # 检查会议室是否可用
        if not room.is_active:
            flash('该会议室当前不可用', 'danger')
            return redirect(url_for('reserve'))
        
        # 检查容量是否足够
        if attendees > room.capacity:
            flash(f'参会人数超过会议室容量（{room.capacity}人）', 'danger')
            return redirect(url_for('reserve'))
        
        # 检查是否有维护计划
        maintenance = Maintenance.query.filter_by(room_id=room_id).filter(
            Maintenance.start_date <= date,
            Maintenance.end_date >= date
        ).first()
        
        if maintenance:
            flash(f'该会议室在所选日期处于维护状态，维护时间：{maintenance.start_date} 至 {maintenance.end_date}', 'danger')
            return redirect(url_for('reserve'))
        
        # 检查时间冲突
        conflicting_reservations = Reservation.query.filter_by(
            room_id=room_id, 
            date=date,
            status='confirmed'
        ).filter(
            ((Reservation.start_time <= start_time) & (Reservation.end_time > start_time)) |
            ((Reservation.start_time < end_time) & (Reservation.end_time >= end_time)) |
            ((Reservation.start_time >= start_time) & (Reservation.end_time <= end_time))
        ).all()
        
        if conflicting_reservations:
            flash('该时间段会议室已被预订，请选择其他时间或会议室', 'danger')
            return redirect(url_for('reserve'))
        
        reservation = Reservation(
            title=title,
            date=date,
            start_time=start_time,
            end_time=end_time,
            attendees=attendees,
            description=description,
            user_id=current_user.id,
            room_id=room_id
        )
        
        db.session.add(reservation)
        db.session.commit()
        flash('会议室预订成功！', 'success')
        return redirect(url_for('my_reservations'))
    
    rooms = Room.query.filter_by(is_active=True).all()
    today = datetime.now().date()
    return render_template('reserve.html', rooms=rooms, today=today)

@app.route('/check_availability', methods=['POST'])
@login_required
def check_availability():
    room_id = request.form.get('room_id')
    date_str = request.form.get('date')
    
    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    
    # 获取会议室
    room = Room.query.get_or_404(room_id)
    
    # 检查是否有维护计划
    maintenance = Maintenance.query.filter_by(room_id=room_id).filter(
        Maintenance.start_date <= date,
        Maintenance.end_date >= date
    ).first()
    
    if maintenance:
        return jsonify({
            'available': False,
            'message': f'该会议室在所选日期处于维护状态，维护时间：{maintenance.start_date} 至 {maintenance.end_date}'
        })
    
    # 获取当天所有预订
    reservations = Reservation.query.filter_by(
        room_id=room_id, 
        date=date,
        status='confirmed'
    ).order_by(Reservation.start_time).all()
    
    # 构建时间段列表
    time_slots = []
    for reservation in reservations:
        time_slots.append({
            'start': reservation.start_time.strftime('%H:%M'),
            'end': reservation.end_time.strftime('%H:%M'),
            'title': reservation.title
        })
    
    return jsonify({
        'available': True,
        'time_slots': time_slots
    })

@app.route('/my_reservations')
@login_required
def my_reservations():
    reservations = Reservation.query.filter_by(user_id=current_user.id).order_by(
        Reservation.date.desc(), 
        Reservation.start_time
    ).all()
    
    # 分类预订
    upcoming = []
    past = []
    canceled = []
    
    for res in reservations:
        if res.status == 'canceled':
            canceled.append(res)
        elif res.is_past():
            past.append(res)
        else:
            upcoming.append(res)
    
    return render_template('my_reservations.html', 
                          upcoming=upcoming, 
                          past=past, 
                          canceled=canceled)

@app.route('/reservation/<int:id>')
@login_required
def reservation_detail(id):
    reservation = Reservation.query.get_or_404(id)
    
    # 检查权限
    if reservation.user_id != current_user.id and not current_user.is_admin():
        flash('您无权查看此预订信息', 'danger')
        return redirect(url_for('my_reservations'))
    
    return render_template('reservation_detail.html', reservation=reservation)

@app.route('/reservation/<int:id>/cancel', methods=['POST'])
@login_required
def cancel_reservation(id):
    reservation = Reservation.query.get_or_404(id)
    
    # 检查权限
    if reservation.user_id != current_user.id and not current_user.is_admin():
        flash('您无权取消此预订', 'danger')
        return redirect(url_for('my_reservations'))
    
    # 检查是否可以取消
    if not reservation.can_cancel():
        flash('无法取消此预订', 'danger')
        return redirect(url_for('my_reservations'))
    
    reservation.status = 'canceled'
    db.session.commit()
    flash('预订已取消', 'success')
    return redirect(url_for('my_reservations'))

# 文档上传路由
@app.route('/reservation/<int:id>/upload', methods=['POST'])
@login_required
def upload_document(id):
    reservation = Reservation.query.get_or_404(id)
    
    # 检查权限
    if reservation.user_id != current_user.id and not current_user.is_admin():
        flash('您无权为此预订上传文档', 'danger')
        return redirect(url_for('my_reservations'))
    
    # 检查预订状态
    if reservation.status != 'confirmed':
        flash('只能为已确认的预订上传文档', 'danger')
        return redirect(url_for('reservation_detail', id=id))
    
    # 检查是否有文件上传
    if 'document' not in request.files:
        flash('未选择任何文件', 'danger')
        return redirect(url_for('reservation_detail', id=id))
    
    file = request.files['document']
    
    # 如果用户未选择文件，浏览器会提交一个空文件
    if file.filename == '':
        flash('未选择任何文件', 'danger')
        return redirect(url_for('reservation_detail', id=id))
    
    if file and allowed_file(file.filename):
        # 安全处理文件名
        original_filename = secure_filename(file.filename)
        
        # 获取文件扩展名（添加错误处理）
        try:
            file_extension = original_filename.rsplit('.', 1)[1].lower()
        except IndexError:
            # 如果没有扩展名，使用空字符串
            file_extension = ""
        
        # 生成唯一的文件名存储（即使没有扩展名）
        if file_extension:
            stored_filename = f"{uuid.uuid4().hex}.{file_extension}"
        else:
            stored_filename = f"{uuid.uuid4().hex}"
        
        # 保存文件
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], stored_filename)
        file.save(file_path)
        
        # 获取文件类型和大小
        file_type = file_extension if file_extension else "unknown"
        file_size = os.path.getsize(file_path)
        
        # 创建文档记录
        document = ReservationDocument(
            filename=original_filename,
            stored_filename=stored_filename,
            file_type=file_type,
            file_size=file_size,
            reservation_id=id,
            uploaded_by=current_user.id
        )
        
        db.session.add(document)
        db.session.commit()
        
        flash('文档上传成功', 'success')
    else:
        flash('不支持的文件格式', 'danger')
    
    return redirect(url_for('reservation_detail', id=id))

# 文档下载路由
@app.route('/document/<int:id>/download')
@login_required
def download_document(id):
    document = ReservationDocument.query.get_or_404(id)
    reservation = document.reservation
    
    # 检查权限
    if reservation.user_id != current_user.id and not current_user.is_admin():
        flash('您无权下载此文档', 'danger')
        return redirect(url_for('my_reservations'))
    
    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        document.stored_filename,
        download_name=document.filename,
        as_attachment=True
    )

# 文档删除路由
@app.route('/document/<int:id>/delete', methods=['POST'])
@login_required
def delete_document(id):
    document = ReservationDocument.query.get_or_404(id)
    reservation = document.reservation
    
    # 检查权限
    if document.uploaded_by != current_user.id and not current_user.is_admin():
        flash('您无权删除此文档', 'danger')
        return redirect(url_for('reservation_detail', id=reservation.id))
    
    # 删除物理文件
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.stored_filename)
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        app.logger.error(f"删除文件时出错: {str(e)}")
    
    # 删除数据库记录
    db.session.delete(document)
    db.session.commit()
    
    flash('文档已删除', 'success')
    return redirect(url_for('reservation_detail', id=reservation.id))

@app.route('/calendar')
def calendar_view():
    today = datetime.now().date()
    year = request.args.get('year', today.year, type=int)
    month = request.args.get('month', today.month, type=int)
    
    # 获取所有会议室
    rooms = Room.query.filter_by(is_active=True).all()
    
    return render_template('calendar.html', 
                          year=year, 
                          month=month,
                          rooms=rooms,
                          today=today)

@app.route('/calendar_data')
def calendar_data():
    year = request.args.get('year', type=int)
    month = request.args.get('month', type=int)
    room_id = request.args.get('room_id', type=int)
    
    # 计算月份的第一天和最后一天
    first_day = datetime(year, month, 1).date()
    if month == 12:
        last_day = datetime(year + 1, 2, 1).date() - timedelta(days=1)
    else:
        last_day = datetime(year, month + 3, 1).date() - timedelta(days=1)
    
    # 查询条件
    query = Reservation.query.filter(
        Reservation.date >= first_day,
        Reservation.date <= last_day,
        Reservation.status == 'confirmed'
    )
    
    if room_id:
        query = query.filter_by(room_id=room_id)
    
    reservations = query.all()
    
    # 构建日历数据
    events = []
    for res in reservations:
        events.append({
            'id': res.id,
            'title': res.title,
            'start': f"{res.date.isoformat()}T{res.start_time.strftime('%H:%M:%S')}",
            'end': f"{res.date.isoformat()}T{res.end_time.strftime('%H:%M:%S')}",
            'room': res.room.name,
            'url': url_for('reservation_detail', id=res.id)
        })
    
    # 获取维护数据
    maintenance_query = Maintenance.query
    if room_id:
        maintenance_query = maintenance_query.filter_by(room_id=room_id)
    
    maintenances = maintenance_query.filter(
        (Maintenance.start_date <= last_day) & (Maintenance.end_date >= first_day)
    ).all()
    
    for m in maintenances:
        events.append({
            'id': f"m{m.id}",
            'title': f"维护: {m.room.name}",
            'start': m.start_date.isoformat(),
            'end': (m.end_date + timedelta(days=1)).isoformat(),  # 全天事件需要+1天
            'color': '#ff9f89',
            'allDay': True
        })
    
    return jsonify(events)

# 管理员路由
@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    # 今日预订统计
    today = datetime.now().date()
    today_reservations = Reservation.query.filter_by(date=today, status='confirmed').count()
    
    # 本周预订统计
    week_start = today - timedelta(days=today.weekday())
    week_end = week_start + timedelta(days=6)
    week_reservations = Reservation.query.filter(
        Reservation.date >= week_start,
        Reservation.date <= week_end,
        Reservation.status == 'confirmed'
    ).count()
    
    # 会议室使用率
    rooms = Room.query.all()
    room_usage = []
    for room in rooms:
        total_reservations = Reservation.query.filter_by(
            room_id=room.id,
            status='confirmed'
        ).filter(
            Reservation.date >= week_start,
            Reservation.date <= week_end
        ).count()
        room_usage.append({
            'name': room.name,
            'count': total_reservations
        })
    
    # 用户统计
    total_users = User.query.count()
    
    return render_template('admin/dashboard.html',
                          today_reservations=today_reservations,
                          week_reservations=week_reservations,
                          room_usage=room_usage,
                          total_users=total_users,
                          rooms=rooms)

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    users = User.query.all()
    return render_template('admin/users.html', users=users)

@app.route('/admin/user/new', methods=['GET', 'POST'])
@login_required
@admin_required
def new_user():
    roles = Role.query.all()
    
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        department = request.form['department']
        phone = request.form['phone']
        role_id = int(request.form['role_id'])
        
        # 检查用户名是否已存在
        if User.query.filter_by(username=username).first():
            flash('用户名已存在', 'danger')
            return render_template('admin/user_form.html', user=None, roles=roles)
        
        # 检查邮箱是否已存在
        if User.query.filter_by(email=email).first():
            flash('邮箱已存在', 'danger')
            return render_template('admin/user_form.html', user=None, roles=roles)
        
        user = User(
            username=username,
            email=email,
            department=department,
            phone=phone,
            role_id=role_id
        )
        user.set_password(password)
        
        db.session.add(user)
        db.session.commit()
        flash('用户添加成功', 'success')
        return redirect(url_for('admin_users'))
    
    return render_template('admin/user_form.html', user=None, roles=roles)

@app.route('/admin/user/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_user(id):
    user = User.query.get_or_404(id)
    roles = Role.query.all()
    
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        department = request.form['department']
        phone = request.form['phone']
        role_id = int(request.form['role_id'])
        new_password = request.form.get('password')
        
        # 检查用户名是否已被其他用户使用
        existing_user = User.query.filter_by(username=username).first()
        if existing_user and existing_user.id != user.id:
            flash('用户名已存在', 'danger')
            return render_template('admin/user_form.html', user=user, roles=roles)
        
        # 检查邮箱是否已被其他用户使用
        existing_email = User.query.filter_by(email=email).first()
        if existing_email and existing_email.id != user.id:
            flash('邮箱已存在', 'danger')
            return render_template('admin/user_form.html', user=user, roles=roles)
        
        # 防止管理员取消自己的管理员权限
        admin_role = Role.query.filter_by(name='admin').first()
        if user.id == current_user.id and user.role_id == admin_role.id and role_id != admin_role.id:
            flash('不能取消自己的管理员权限', 'danger')
            return render_template('admin/user_form.html', user=user, roles=roles)
        
        user.username = username
        user.email = email
        user.department = department
        user.phone = phone
        user.role_id = role_id
        
        # 如果提供了新密码，则更新密码
        if new_password:
            user.set_password(new_password)
        
        db.session.commit()
        flash('用户信息已更新', 'success')
        return redirect(url_for('admin_users'))
    
    return render_template('admin/user_form.html', user=user, roles=roles)

@app.route('/admin/user/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_user(id):
    user = User.query.get_or_404(id)
    
    # 防止删除自己
    if user.id == current_user.id:
        flash('不能删除自己的账户', 'danger')
        return redirect(url_for('admin_users'))
    
    # 删除用户关联的预订记录
    Reservation.query.filter_by(user_id=user.id).delete()
    
    # 删除用户
    db.session.delete(user)
    db.session.commit()
    flash('用户已删除', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:id>/toggle_admin', methods=['POST'])
@login_required
@admin_required
def toggle_admin(id):
    user = User.query.get_or_404(id)
    
    # 防止自己取消自己的管理员权限
    if user.id == current_user.id:
        flash('不能修改自己的权限', 'danger')
        return redirect(url_for('admin_users'))
    
    admin_role = Role.query.filter_by(name='admin').first()
    user_role = Role.query.filter_by(name='user').first()
    
    if user.role_id == admin_role.id:
        user.role_id = user_role.id
        flash(f'已移除 {user.username} 的管理员权限', 'success')
    else:
        user.role_id = admin_role.id
        flash(f'已将 {user.username} 设为管理员', 'success')
    
    db.session.commit()
    return redirect(url_for('admin_users'))

@app.route('/admin/rooms')
@login_required
@admin_required
def admin_rooms():
    rooms = Room.query.all()
    return render_template('admin/rooms.html', rooms=rooms)

@app.route('/admin/room/new', methods=['GET', 'POST'])
@login_required
@admin_required
def new_room():
    equipments = Equipment.query.all()
    
    if request.method == 'POST':
        name = request.form['name']
        location = request.form['location']
        capacity = int(request.form['capacity'])
        description = request.form['description']
        
        room = Room(
            name=name,
            location=location,
            capacity=capacity,
            description=description,
            is_active=True
        )
        db.session.add(room)
        db.session.flush()  # 获取room.id
        
        # 处理设备
        for equipment in equipments:
            if f'equipment_{equipment.id}' in request.form:
                quantity = int(request.form.get(f'quantity_{equipment.id}', 1))
                room_equipment = RoomEquipment(
                    room_id=room.id,
                    equipment_id=equipment.id,
                    quantity=quantity
                )
                db.session.add(room_equipment)
        
        db.session.commit()
        flash('会议室添加成功', 'success')
        return redirect(url_for('admin_rooms'))
    
    return render_template('admin/room_form.html', 
                          room=None, 
                          equipments=equipments,
                          room_equipments=[])

@app.route('/admin/room/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_room(id):
    room = Room.query.get_or_404(id)
    equipments = Equipment.query.all()
    room_equipments = {re.equipment_id: re.quantity for re in room.equipments}
    
    if request.method == 'POST':
        room.name = request.form['name']
        room.location = request.form['location']
        room.capacity = int(request.form['capacity'])
        room.description = request.form['description']
        room.is_active = 'is_active' in request.form
        
        # 清除原有设备关联
        RoomEquipment.query.filter_by(room_id=room.id).delete()
        
        # 重新添加设备关联
        for equipment in equipments:
            if f'equipment_{equipment.id}' in request.form:
                quantity = int(request.form.get(f'quantity_{equipment.id}', 1))
                room_equipment = RoomEquipment(
                    room_id=room.id,
                    equipment_id=equipment.id,
                    quantity=quantity
                )
                db.session.add(room_equipment)
        
        db.session.commit()
        flash('会议室信息已更新', 'success')
        return redirect(url_for('admin_rooms'))
    
    return render_template('admin/room_form.html', 
                          room=room, 
                          equipments=equipments,
                          room_equipments=room_equipments)

@app.route('/admin/room/<int:id>/toggle_status', methods=['POST'])
@login_required
@admin_required
def toggle_room_status(id):
    room = Room.query.get_or_404(id)
    room.is_active = not room.is_active
    
    status_text = "启用" if room.is_active else "禁用"
    db.session.commit()
    flash(f'会议室 {room.name} 已{status_text}', 'success')
    return redirect(url_for('admin_rooms'))

@app.route('/admin/equipment')
@login_required
@admin_required
def admin_equipment():
    equipments = Equipment.query.all()
    return render_template('admin/equipment.html', equipments=equipments)

@app.route('/admin/equipment/new', methods=['GET', 'POST'])
@login_required
@admin_required
def new_equipment():
    if request.method == 'POST':
        name = request.form['name']
        
        if Equipment.query.filter_by(name=name).first():
            flash('设备名称已存在', 'danger')
            return redirect(url_for('new_equipment'))
        
        equipment = Equipment(name=name)
        db.session.add(equipment)
        db.session.commit()
        flash('设备添加成功', 'success')
        return redirect(url_for('admin_equipment'))
    
    return render_template('admin/equipment_form.html', equipment=None)

@app.route('/admin/equipment/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_equipment(id):
    equipment = Equipment.query.get_or_404(id)
    
    if request.method == 'POST':
        name = request.form['name']
        
        if name != equipment.name and Equipment.query.filter_by(name=name).first():
            flash('设备名称已存在', 'danger')
            return redirect(url_for('edit_equipment', id=id))
        
        equipment.name = name
        db.session.commit()
        flash('设备信息已更新', 'success')
        return redirect(url_for('admin_equipment'))
    
    return render_template('admin/equipment_form.html', equipment=equipment)

@app.route('/admin/equipment/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_equipment(id):
    equipment = Equipment.query.get_or_404(id)
    
    # 检查是否有关联的会议室
    if RoomEquipment.query.filter_by(equipment_id=id).first():
        flash('该设备已关联到会议室，无法删除', 'danger')
        return redirect(url_for('admin_equipment'))
    
    db.session.delete(equipment)
    db.session.commit()
    flash('设备已删除', 'success')
    return redirect(url_for('admin_equipment'))

@app.route('/admin/maintenance')
@login_required
@admin_required
def admin_maintenance():
    maintenances = Maintenance.query.order_by(Maintenance.start_date.desc()).all()
    return render_template('admin/maintenance.html', maintenances=maintenances)

@app.route('/admin/maintenance/new', methods=['GET', 'POST'])
@login_required
@admin_required
def new_maintenance():
    rooms = Room.query.all()
    
    if request.method == 'POST':
        room_id = request.form['room_id']
        start_date_str = request.form['start_date']
        end_date_str = request.form['end_date']
        reason = request.form['reason']
        
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        
        # 验证日期
        if start_date > end_date:
            flash('结束日期必须晚于开始日期', 'danger')
            return redirect(url_for('new_maintenance'))
        
        # 检查是否有冲突的预订
        conflicting_reservations = Reservation.query.filter_by(
            room_id=room_id,
            status='confirmed'
        ).filter(
            Reservation.date >= start_date,
            Reservation.date <= end_date
        ).all()
        
        if conflicting_reservations:
            flash('该时间段内有已确认的预订，请先取消这些预订', 'danger')
            return redirect(url_for('new_maintenance'))
        
        maintenance = Maintenance(
            room_id=room_id,
            start_date=start_date,
            end_date=end_date,
            reason=reason,
            created_by=current_user.id
        )
        
        db.session.add(maintenance)
        db.session.commit()
        flash('维护计划已添加', 'success')
        return redirect(url_for('admin_maintenance'))
    
    return render_template('admin/maintenance_form.html', rooms=rooms, maintenance=None)

@app.route('/admin/maintenance/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_maintenance(id):
    maintenance = Maintenance.query.get_or_404(id)
    rooms = Room.query.all()
    
    if request.method == 'POST':
        room_id = request.form['room_id']
        start_date_str = request.form['start_date']
        end_date_str = request.form['end_date']
        reason = request.form['reason']
        
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        
        # 验证日期
        if start_date > end_date:
            flash('结束日期必须晚于开始日期', 'danger')
            return redirect(url_for('edit_maintenance', id=id))
        
        # 如果修改了会议室或日期，检查冲突
        if (room_id != str(maintenance.room_id) or 
            start_date != maintenance.start_date or 
            end_date != maintenance.end_date):
            
            conflicting_reservations = Reservation.query.filter_by(
                room_id=room_id,
                status='confirmed'
            ).filter(
                Reservation.date >= start_date,
                Reservation.date <= end_date
            ).all()
            
            if conflicting_reservations:
                flash('该时间段内有已确认的预订，请先取消这些预订', 'danger')
                return redirect(url_for('edit_maintenance', id=id))
        
        maintenance.room_id = room_id
        maintenance.start_date = start_date
        maintenance.end_date = end_date
        maintenance.reason = reason
        
        db.session.commit()
        flash('维护计划已更新', 'success')
        return redirect(url_for('admin_maintenance'))
    
    return render_template('admin/maintenance_form.html', 
                          rooms=rooms, 
                          maintenance=maintenance)

@app.route('/admin/maintenance/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_maintenance(id):
    maintenance = Maintenance.query.get_or_404(id)
    db.session.delete(maintenance)
    db.session.commit()
    flash('维护计划已删除', 'success')
    return redirect(url_for('admin_maintenance'))

@app.route('/admin/reservations')
@login_required
@admin_required
def admin_reservations():
    status = request.args.get('status', 'all')
    room_id = request.args.get('room_id', 'all')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    
    query = Reservation.query
    
    # 筛选状态
    if status != 'all':
        query = query.filter_by(status=status)
    
    # 筛选会议室
    if room_id != 'all' and room_id.isdigit():
        query = query.filter_by(room_id=int(room_id))
    
    # 筛选日期范围
    if date_from:
        from_date = datetime.strptime(date_from, '%Y-%m-%d').date()
        query = query.filter(Reservation.date >= from_date)
    
    if date_to:
        to_date = datetime.strptime(date_to, '%Y-%m-%d').date()
        query = query.filter(Reservation.date <= to_date)
    
    reservations = query.order_by(Reservation.date.desc(), Reservation.start_time).all()
    rooms = Room.query.all()
    
    return render_template('admin/reservations.html', 
                          reservations=reservations,
                          rooms=rooms,
                          status=status,
                          room_id=room_id,
                          date_from=date_from,
                          date_to=date_to)

@app.route('/admin/reservation/<int:id>/cancel', methods=['POST'])
@login_required
@admin_required
def admin_cancel_reservation(id):
    reservation = Reservation.query.get_or_404(id)
    
    if reservation.status != 'confirmed':
        flash('只能取消已确认的预订', 'danger')
        return redirect(url_for('admin_reservations'))
    
    reservation.status = 'canceled'
    db.session.commit()
    
    # 发送通知给用户（这里简化处理，实际应该发送邮件或其他通知）
    flash(f'预订 {reservation.title} 已被取消', 'success')
    return redirect(url_for('admin_reservations'))

@app.route('/admin/reports')
@login_required
@admin_required
def admin_reports():
    # 获取报表类型
    report_type = request.args.get('type', 'room_usage')
    
    # 获取时间范围
    date_from_str = request.args.get('date_from', '')
    date_to_str = request.args.get('date_to', '')
    
    today = datetime.now().date()
    
    # 默认为过去30天
    if not date_from_str:
        date_from = today - timedelta(days=30)
        date_from_str = date_from.strftime('%Y-%m-%d')
    else:
        date_from = datetime.strptime(date_from_str, '%Y-%m-%d').date()
    
    if not date_to_str:
        date_to = today
        date_to_str = date_to.strftime('%Y-%m-%d')
    else:
        date_to = datetime.strptime(date_to_str, '%Y-%m-%d').date()
    
    # 获取所有会议室
    rooms = Room.query.all()
    
    # 生成报表数据
    if report_type == 'room_usage':
        # 会议室使用率报表
        report_data = []
        
        for room in rooms:
            # 计算总预订数
            total_reservations = Reservation.query.filter_by(
                room_id=room.id,
                status='confirmed'
            ).filter(
                Reservation.date >= date_from,
                Reservation.date <= date_to
            ).count()
            
            # 计算总使用时长（小时）
            total_hours = 0
            reservations = Reservation.query.filter_by(
                room_id=room.id,
                status='confirmed'
            ).filter(
                Reservation.date >= date_from,
                Reservation.date <= date_to
            ).all()
            
            for res in reservations:
                start_datetime = datetime.combine(res.date, res.start_time)
                end_datetime = datetime.combine(res.date, res.end_time)
                duration = (end_datetime - start_datetime).total_seconds() / 3600  # 转换为小时
                total_hours += duration
            
            report_data.append({
                'room_name': room.name,
                'total_reservations': total_reservations,
                'total_hours': round(total_hours, 1)
            })
        
        return render_template('admin/reports.html',
                              report_type=report_type,
                              date_from=date_from_str,
                              date_to=date_to_str,
                              report_data=report_data)
    
    elif report_type == 'user_activity':
        # 用户活动报表
        users = User.query.all()
        report_data = []
        
        for user in users:
            # 计算用户预订数
            total_reservations = Reservation.query.filter_by(
                user_id=user.id
            ).filter(
                Reservation.date >= date_from,
                Reservation.date <= date_to
            ).count()
            
            # 计算取消预订数
            canceled_reservations = Reservation.query.filter_by(
                user_id=user.id,
                status='canceled'
            ).filter(
                Reservation.date >= date_from,
                Reservation.date <= date_to
            ).count()
            
            if total_reservations > 0:  # 只显示有活动的用户
                report_data.append({
                    'username': user.username,
                    'department': user.department,
                    'total_reservations': total_reservations,
                    'canceled_reservations': canceled_reservations
                })
        
        return render_template('admin/reports.html',
                              report_type=report_type,
                              date_from=date_from_str,
                              date_to=date_to_str,
                report_data=report_data)
    
    elif report_type == 'time_distribution':
        # 时间分布报表
        # 按小时统计预订数量
        hours = list(range(8, 21))  # 假设工作时间为8:00-20:00
        hour_counts = {hour: 0 for hour in hours}
        
        # 按星期几统计预订数量
        weekdays = ['周一', '周二', '周三', '周四', '周五', '周六', '周日']
        weekday_counts = {i: 0 for i in range(7)}
        
        reservations = Reservation.query.filter_by(
            status='confirmed'
        ).filter(
            Reservation.date >= date_from,
            Reservation.date <= date_to
        ).all()
        
        for res in reservations:
            # 统计每个小时的预订数
            start_hour = res.start_time.hour
            end_hour = res.end_time.hour
            
            for hour in range(start_hour, end_hour):
                if hour in hour_counts:
                    hour_counts[hour] += 1
            
            # 统计每个工作日的预订数
            weekday = res.date.weekday()  # 0=周一, 6=周日
            weekday_counts[weekday] += 1
        
        hour_data = [{'hour': f"{hour}:00", 'count': count} for hour, count in hour_counts.items()]
        weekday_data = [{'weekday': weekdays[i], 'count': count} for i, count in weekday_counts.items()]
        
        return render_template('admin/reports.html',
                              report_type=report_type,
                              date_from=date_from_str,
                              date_to=date_to_str,
                              hour_data=hour_data,
                              weekday_data=weekday_data)
    
    return render_template('admin/reports.html',
                          report_type=report_type,
                          date_from=date_from_str,
                          date_to=date_to_str)

@app.route('/api/export_reservations', methods=['GET'])
@login_required
@admin_required
def export_reservations():
    from io import StringIO
    import csv
    
    status = request.args.get('status', 'all')
    room_id = request.args.get('room_id', 'all')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    
    query = Reservation.query
    
    # 筛选状态
    if status != 'all':
        query = query.filter_by(status=status)
    
    # 筛选会议室
    if room_id != 'all' and room_id.isdigit():
        query = query.filter_by(room_id=int(room_id))
    
    # 筛选日期范围
    if date_from:
        from_date = datetime.strptime(date_from, '%Y-%m-%d').date()
        query = query.filter(Reservation.date >= from_date)
    
    if date_to:
        to_date = datetime.strptime(date_to, '%Y-%m-%d').date()
        query = query.filter(Reservation.date <= to_date)
    
    reservations = query.order_by(Reservation.date, Reservation.start_time).all()
    
    # 创建CSV
    si = StringIO()
    csv_writer = csv.writer(si)
    
    # 写入表头
    csv_writer.writerow([
        '预订ID', '标题', '日期', '开始时间', '结束时间', 
        '参会人数', '会议室', '预订人', '部门', '状态', '创建时间'
    ])
    
    # 写入数据
    for res in reservations:
        csv_writer.writerow([
            res.id,
            res.title,
            res.date.strftime('%Y-%m-%d'),
            res.start_time.strftime('%H:%M'),
            res.end_time.strftime('%H:%M'),
            res.attendees,
            res.room.name,
            res.user.username,
            res.user.department,
            res.status,
            res.created_at.strftime('%Y-%m-%d %H:%M:%S')
        ])
    
    output = si.getvalue()
    
    # 设置响应头
    from flask import Response
    filename = f"reservations_{datetime.now().strftime('%Y%m%d')}.csv"
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment;filename={filename}"}
    )

# 错误处理
@app.errorhandler(404)
def page_not_found(e):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('errors/500.html'), 500

# API路由
@app.route('/api/rooms')
def api_rooms():
    rooms = Room.query.filter_by(is_active=True).all()
    room_list = []
    
    for room in rooms:
        equipment_list = [{'id': eq.equipment.id, 'name': eq.equipment.name, 'quantity': eq.quantity} 
                         for eq in room.equipments]
        
        room_list.append({
            'id': room.id,
            'name': room.name,
            'location': room.location,
            'capacity': room.capacity,
            'description': room.description,
            'equipment': equipment_list
        })
    
    return jsonify(room_list)

@app.route('/api/room_availability')
def api_room_availability():
    room_id = request.args.get('room_id', type=int)
    date_str = request.args.get('date')
    
    if not room_id or not date_str:
        return jsonify({'error': '缺少必要参数'}), 400
    
    try:
        date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': '日期格式无效'}), 400
    
    # 获取会议室
    room = Room.query.get_or_404(room_id)
    
    # 检查会议室是否可用
    if not room.is_active:
        return jsonify({
            'available': False,
            'reason': '会议室当前不可用'
        })
    
    # 检查是否有维护计划
    maintenance = Maintenance.query.filter_by(room_id=room_id).filter(
        Maintenance.start_date <= date,
        Maintenance.end_date >= date
    ).first()
    
    if maintenance:
        return jsonify({
            'available': False,
            'reason': f'维护中 ({maintenance.start_date} 至 {maintenance.end_date})',
            'maintenance_id': maintenance.id
        })
    
    # 获取当天所有预订
    reservations = Reservation.query.filter_by(
        room_id=room_id, 
        date=date,
        status='confirmed'
    ).order_by(Reservation.start_time).all()
    
    # 工作时间 (8:00-20:00)
    work_start = time(8, 0)
    work_end = time(20, 0)
    
    # 计算可用时间段
    available_slots = []
    current_time = work_start
    
    for res in reservations:
        # 如果当前时间早于预订开始时间，添加一个可用时间段
        if current_time < res.start_time:
            available_slots.append({
                'start': current_time.strftime('%H:%M'),
                'end': res.start_time.strftime('%H:%M')
            })
        
        # 更新当前时间为预订结束时间
        if current_time < res.end_time:
            current_time = res.end_time
    
    # 如果当前时间早于工作结束时间，添加最后一个可用时间段
    if current_time < work_end:
        available_slots.append({
            'start': current_time.strftime('%H:%M'),
            'end': work_end.strftime('%H:%M')
        })
    
    # 构建预订列表
    reservation_list = []
    for res in reservations:
        reservation_list.append({
            'id': res.id,
            'title': res.title,
            'start': res.start_time.strftime('%H:%M'),
            'end': res.end_time.strftime('%H:%M'),
            'user': res.user.username
        })
    
    return jsonify({
        'available': True,
        'available_slots': available_slots,
        'reservations': reservation_list
    })

@app.route('/api/quick_reserve', methods=['POST'])
@login_required
def api_quick_reserve():
    data = request.json
    
    try:
        room_id = data.get('room_id')
        date_str = data.get('date')
        start_time_str = data.get('start_time')
        end_time_str = data.get('end_time')
        title = data.get('title')
        attendees = data.get('attendees', 1)
        description = data.get('description', '')
        
        # 验证必要字段
        if not all([room_id, date_str, start_time_str, end_time_str, title]):
            return jsonify({'error': '缺少必要参数'}), 400
        
        # 转换日期和时间
        date = datetime.strptime(date_str, '%Y-%m-%d').date()
        start_time = datetime.strptime(start_time_str, '%H:%M').time()
        end_time = datetime.strptime(end_time_str, '%H:%M').time()
        
        # 检查日期是否有效
        if date < datetime.now().date():
            return jsonify({'error': '无法预订过去的日期'}), 400
        
        # 检查时间是否有效
        if start_time >= end_time:
            return jsonify({'error': '结束时间必须晚于开始时间'}), 400
        
        # 检查工作时间
        work_start = time(8, 0)
        work_end = time(20, 0)
        if start_time < work_start or end_time > work_end:
            return jsonify({'error': '预订时间必须在工作时间内（8:00-20:00）'}), 400
        
        # 获取会议室
        room = Room.query.get_or_404(room_id)
        
        # 检查会议室是否可用
        if not room.is_active:
            return jsonify({'error': '该会议室当前不可用'}), 400
        
        # 检查容量是否足够
        if attendees > room.capacity:
            return jsonify({'error': f'参会人数超过会议室容量（{room.capacity}人）'}), 400
        
        # 检查是否有维护计划
        maintenance = Maintenance.query.filter_by(room_id=room_id).filter(
            Maintenance.start_date <= date,
            Maintenance.end_date >= date
        ).first()
        
        if maintenance:
            return jsonify({'error': f'该会议室在所选日期处于维护状态'}), 400
        
        # 检查时间冲突
        conflicting_reservations = Reservation.query.filter_by(
            room_id=room_id, 
            date=date,
            status='confirmed'
        ).filter(
            ((Reservation.start_time <= start_time) & (Reservation.end_time > start_time)) |
            ((Reservation.start_time < end_time) & (Reservation.end_time >= end_time)) |
            ((Reservation.start_time >= start_time) & (Reservation.end_time <= end_time))
        ).all()
        
        if conflicting_reservations:
            return jsonify({'error': '该时间段会议室已被预订'}), 400
        
        # 创建预订
        reservation = Reservation(
            title=title,
            date=date,
            start_time=start_time,
            end_time=end_time,
            attendees=attendees,
            description=description,
            user_id=current_user.id,
            room_id=room_id
        )
        
        db.session.add(reservation)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': '会议室预订成功',
            'reservation_id': reservation.id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'预订失败: {str(e)}'}), 500

# 辅助函数
def get_next_week_dates():
    """获取未来一周的日期列表"""
    today = datetime.now().date()
    return [today + timedelta(days=i) for i in range(7)]

app.jinja_env.globals.update(get_next_week_dates=get_next_week_dates)

@app.route('/api/recent_reservations')
@login_required
@admin_required
def api_recent_reservations():
    # 获取最近10条预订记录
    recent_reservations = Reservation.query.order_by(
        Reservation.date.desc(), 
        Reservation.start_time.desc()
    ).limit(10).all()
    
    result = []
    for res in recent_reservations:
        result.append({
            'id': res.id,
            'title': res.title,
            'date': res.date.strftime('%Y-%m-%d'),
            'start_time': res.start_time.strftime('%H:%M'),
            'end_time': res.end_time.strftime('%H:%M'),
            'room_name': res.room.name,
            'username': res.user.username,
            'status': res.status
        })
    
    return jsonify(result)

# 签到表下载路由
@app.route('/reservation/<int:id>/signin-sheet')
@login_required
def download_signin_sheet(id):
    # 获取预订信息
    reservation = Reservation.query.get_or_404(id)
    
    # 检查权限
    if reservation.user_id != current_user.id and not current_user.is_admin():
        flash('您无权下载此签到表', 'danger')
        return redirect(url_for('my_reservations'))
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run('会议签到表')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加会议信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info.add_run(f'会议名称：{reservation.title}\n').font.bold = True
    info.add_run(f'会议时间：{reservation.date.strftime("%Y年%m月%d日")} {reservation.format_time()}\n').font.bold = True
    info.add_run(f'会议地点：{reservation.room.name} ({reservation.room.location})\n').font.bold = True
    
    # 添加说明
    doc.add_paragraph('请与会人员签到：', style='Intense Quote')
    
    # 创建签到表格
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '序号'
    header_cells[1].text = '部门'
    header_cells[2].text = '姓名'
    header_cells[3].text = '签到'
    
    # 设置表头格式
    for i in range(4):
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
    
    # 填充序号
    for i in range(1, 11):
        cell = table.cell(i, 0)
        cell.text = str(i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 设置列宽
    table.columns[0].width = Cm(1.5)  # 序号列
    table.columns[1].width = Cm(5)    # 部门列
    table.columns[2].width = Cm(4)    # 姓名列
    table.columns[3].width = Cm(4)    # 签到列
    
    # 添加页脚
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"本文档由会议室预订系统自动生成 - {datetime.now().strftime('%Y-%m-%d')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存到内存
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # 设置文件名
    filename = f"签到表_{reservation.title}_{reservation.date.strftime('%Y%m%d')}.docx"
    
    # 返回文件
    return send_file(
        file_stream, 
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        download_name=filename,
        as_attachment=True
    )

if __name__ == '__main__':
    # app.run(debug=True)
    # ### 默认部署方式：开发模式运行，基于 Flask 的内置调试服务器（默认端口 5000）
    app.run(host='10.31.2.9', port=5001, debug=True)
